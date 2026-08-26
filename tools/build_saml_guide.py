from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path("docs/Cebu_Rooms_SAML_2_0_Implementation_Guide.docx")
BLUE = "2E74B5"
DARK = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa)-1)]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def font(run, size=11, bold=None, color=INK, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def add_p(doc, text="", bold_label=None, style=None, after=6, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_label and text.startswith(bold_label):
        font(p.add_run(bold_label), bold=True)
        font(p.add_run(text[len(bold_label):]))
    else:
        font(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text))
    return p


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(.18)
        p.paragraph_format.right_indent = Inches(.12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2F4F7")
        pPr.append(shd)
        font(p.add_run(line if line else " "), size=8.5, color="1F2937", name="Courier New")
    add_p(doc, "", after=4)


def add_callout(doc, label, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.18)
    p.paragraph_format.right_indent = Inches(.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE)
    pPr.append(shd)
    r = font(p.add_run(label.upper() + "  "), bold=True, color=color)
    font(p.add_run(text), color=INK)
    return p


def add_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(h), size=9.5, bold=True, color=DARK)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            font(p.add_run(str(value)), size=9.2, color=INK)
    set_table_geometry(t, widths)
    for row in t.rows:
        row._tr.get_or_add_trPr()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(.492)
sec.footer_distance = Inches(.492)

# Compact reference guide preset tokens.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True
for name in ["List Bullet", "List Bullet 2", "List Number"]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.25

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(header.add_run("CEBU ROOMS  |  SAML 2.0 IMPLEMENTATION GUIDE"), size=8.5, bold=True, color=MUTED)
footer = sec.footer.paragraphs[0]
add_page_field(footer)

# Editorial cover.
add_p(doc, "IMPLEMENTATION PLAYBOOK", after=18).alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("SAML 2.0 for Cebu Rooms"), size=28, bold=True, color=DARK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Development plan, setup, configuration, testing, deployment, and replication runbook"), size=14, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
font(p.add_run("Target stack: Laravel 12 · PHP 8.2 · Inertia/Vue · OneLogin PHP-SAML 4.3"), size=10.5, bold=True, color=INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Prepared from the implementation in cebu_rooms2 · 15 July 2026"), size=10, color=MUTED)
doc.add_paragraph().paragraph_format.space_after = Pt(90)
add_callout(doc, "Purpose", "Use this guide to reproduce the current service-provider-initiated SAML login flow in another environment or Laravel application, while preserving the same user-matching and security controls.")
doc.add_page_break()

doc.add_heading("How to use this guide", level=1)
add_p(doc, "This is both a delivery plan and an operator runbook. Sections 1–4 explain the target behavior and development sequence. Sections 5–11 provide the exact setup, partner exchange, verification, release, and troubleshooting process.")
add_callout(doc, "Current implementation boundary", "Cebu Rooms acts as the SAML Service Provider (SP). The external AMS/OnePortal system acts as the Identity Provider (IdP). Login is SP-initiated. The current /saml2/logout endpoint ends the local Rooms flow but does not yet issue a standards-based SAML LogoutRequest to the IdP.", GOLD)
doc.add_heading("Document map", level=2)
for item in [
    "1. Target outcome and architecture",
    "2. Development plan and acceptance gates",
    "3. Inputs, ownership, and readiness",
    "4. Implementation blueprint",
    "5. Application and database setup",
    "6. Environment and SAML configuration",
    "7. IdP registration and attribute contract",
    "8. End-to-end test procedure",
    "9. Deployment, rollback, and operations",
    "10. Troubleshooting",
    "11. Replication checklist and appendices",
]: add_bullet(doc, item)

doc.add_heading("1. Target outcome and architecture", level=1)
doc.add_heading("1.1 Desired behavior", level=2)
add_p(doc, "A user selects SAML login in Rooms. Rooms generates an AuthnRequest and redirects the browser to the active IdP. The IdP authenticates the user and POSTs a signed SAMLResponse to the Assertion Consumer Service (ACS). Rooms validates the response, finds an existing active user account by email, establishes the Laravel session, records an audit event, and redirects to /MainDashboard.")
doc.add_heading("1.2 Trust flow", level=2)
for text in [
    "Browser → GET /saml2/login: Rooms selects the one active IdP configuration and stores the generated request ID with an expiry.",
    "Rooms → IdP SSO URL: an HTTP-Redirect AuthnRequest requests emailAddress NameID and carries a relative RelayState.",
    "IdP → POST /saml2/acs: the browser submits SAMLResponse and RelayState by HTTP-POST.",
    "ACS validation: signature/library validation plus issuer, destination, recipient, audience, time window, InResponseTo, and replay checks.",
    "Account resolution: NameID is preferred; email may fall back to email, mail, EmailAddress, emailAddress, or the standard LDAP email OID attribute.",
    "Session creation: only an existing user_accounts record with account_status=active is admitted; the session is regenerated and last-login fields are updated.",
]: add_number(doc, text)
doc.add_heading("1.3 Endpoint contract", level=2)
add_table(doc, ["Endpoint", "Method / binding", "Purpose"], [
    ("/saml2/metadata", "GET", "Publishes SP Entity ID, ACS, SLO, NameID format, and signing preferences."),
    ("/saml2/login", "GET / HTTP-Redirect", "Creates AuthnRequest and redirects to the active IdP."),
    ("/saml2/acs", "POST / HTTP-POST", "Consumes and validates SAMLResponse; GET restarts login."),
    ("/saml2/logout", "GET or POST", "Records local logout completion and returns to login."),
    ("/saml2/user-not-found", "GET", "Explains missing or inactive local account."),
], [2200, 2100, 5060])

doc.add_heading("2. Development plan and acceptance gates", level=1)
add_table(doc, ["Phase", "Core work", "Exit / acceptance gate"], [
    ("0 · Discovery", "Confirm IdP metadata, domains, proxy/TLS, user source, NameID, signing rules, and owners.", "All partner values documented; test identities and maintenance window approved."),
    ("1 · Foundation", "Install OneLogin PHP-SAML; add services.saml config; create configuration, replay, and audit tables/models.", "Migrations pass; secrets stay outside source control; one active IdP constraint works."),
    ("2 · SP endpoints", "Build metadata, AuthnRequest redirect, ACS, local logout, routes, and admin configuration UI.", "Metadata validates; IdP can import SP details; permissions restrict SAML administration."),
    ("3 · Security", "Enforce strict validation, request correlation, replay defense, clock windows, safe RelayState, session regeneration, and audit logging.", "Negative tests reject altered, expired, replayed, wrong-audience, and wrong-destination responses."),
    ("4 · Federation", "Exchange metadata/certificates, configure AMS/OnePortal, provision test users, and run end-to-end tests.", "Known active user succeeds; unknown/inactive users fail safely; audit evidence exists."),
    ("5 · Release", "Back up, migrate, seed/configure, clear/rebuild caches, smoke test, monitor, and document rollback.", "Production SAML login succeeds and local fallback remains available during stabilization."),
    ("6 · Hardening", "Implement real SLO if required, SP request signing/key separation, certificate rollover, replay cleanup, alerting, and periodic federation tests.", "Security review closes documented gaps and operations accepts the support runbook."),
], [1150, 4210, 4000])
doc.add_heading("2.1 Recommended responsibilities", level=2)
add_table(doc, ["Owner", "Accountability"], [
    ("Rooms application team", "SP code, deployment, APP_URL/proxy correctness, database, logging, and user-account matching."),
    ("AMS/OnePortal administrator", "IdP metadata, signing certificate, SSO/SLO URLs, NameID/attributes, SP registration, and test identities."),
    ("Security / infrastructure", "TLS, secrets, certificate lifecycle, time synchronization, firewall/proxy, log retention, and change approval."),
    ("Business administrator", "Pre-provisioned Rooms accounts, active status, role/permissions, and user acceptance testing."),
], [2250, 7110])

doc.add_heading("3. Inputs, ownership, and readiness", level=1)
doc.add_heading("3.1 Obtain from the IdP administrator", level=2)
for x in [
    "IdP Entity ID (exact, case-sensitive string).",
    "Single Sign-On URL and supported binding; current Rooms flow expects HTTP-Redirect for AuthnRequest.",
    "Single Logout URL, if available; note that full federated logout is a future enhancement in the current code.",
    "Current X.509 signing certificate, fingerprint, expiry date, and rollover plan.",
    "Whether the IdP signs the Response, the Assertion, or both; whether it requires signed AuthnRequests.",
    "NameID format and attribute names. The minimum identity contract is a stable email address.",
    "Two active test accounts and one negative-test identity not provisioned in Rooms.",
]: add_bullet(doc, x)
doc.add_heading("3.2 Provide to the IdP administrator", level=2)
add_code(doc, """SP Entity ID:  ${APP_URL}/saml2/metadata
ACS URL:       ${APP_URL}/saml2/acs
SLO URL:       ${APP_URL}/saml2/logout
Metadata URL:  ${APP_URL}/saml2/metadata
ACS binding:   urn:oasis:names:tc:SAML:2.0:bindings:HTTP-POST
NameID format: urn:oasis:names:tc:SAML:1.1:nameid-format:emailAddress""")
doc.add_heading("3.3 Platform readiness", level=2)
for x in [
    "Public HTTPS URL is final and APP_URL exactly matches it; do not mix hostnames, schemes, ports, or trailing proxies.",
    "Reverse proxy forwards scheme/host correctly and Laravel trusted-proxy behavior produces public HTTPS URLs.",
    "Application and IdP hosts use reliable NTP time synchronization.",
    "Database backup, migration permissions, PHP 8.2+, Composer, and required PHP XML/OpenSSL extensions are available.",
    "At least one non-SAML administrative access path is retained for recovery during rollout.",
]: add_bullet(doc, x)

doc.add_heading("4. Implementation blueprint", level=1)
doc.add_heading("4.1 Repository components", level=2)
add_table(doc, ["Component", "Responsibility"], [
    ("composer.json", "Requires onelogin/php-saml ^4.3."),
    ("config/services.php", "Maps environment variables to services.saml settings."),
    ("SamlConfiguration + migration", "Stores partner metadata and activation/status controls."),
    ("SamlReplayRecord + migration", "Stores request, response, and assertion IDs with expiry; unique indexes prevent reuse."),
    ("SamlAuditEvent + migration", "Records issued requests, metadata views, accepted/rejected assertions, and logout events."),
    ("SamlMetadataController", "Generates SP metadata XML."),
    ("SamlSpController", "Builds AuthnRequest, validates response, resolves account, logs in, and audits."),
    ("SamlConfigurationController", "Admin-only CRUD and metadata XML parsing; enforces one active provider per mode."),
    ("resources/js/Pages/SamlIntegration.vue", "Admin setup and configuration interface."),
    ("tests/Feature/SamlIntegrationTest.php", "Metadata, authorization, activation, ACS success, unknown user, and request correlation tests."),
], [3550, 5810])
doc.add_heading("4.2 Validation sequence", level=2)
for x in [
    "Require SAMLResponse and an active IdP certificate.",
    "Base64-decode and parse XML with network access disabled (LIBXML_NONET).",
    "If InResponseTo is present, require a matching unexpired issued request record.",
    "Run OneLogin validation in strict mode with XML validation and strict destination matching.",
    "Require Success status and exact configured issuer.",
    "Require destination and subject recipient to match the ACS URL.",
    "Require audience to equal SAML_SP_ENTITY_ID.",
    "Require email identity; validate NotBefore and NotOnOrAfter with configured skew.",
    "Reject previously stored response_id or assertion_id, then persist them before login.",
    "Resolve an existing active account case-insensitively by email; regenerate session after Auth::login.",
]: add_number(doc, x)
add_callout(doc, "Security review item", "The database flags require_signed_requests, sign_responses, and sign_assertions are federation metadata/configuration fields, but the current AuthnRequest builder does not sign outbound requests. The current metadata certificate is sourced from SAML_IDP_PUBLIC_CERT, whose name implies an IdP certificate. Before enabling AuthnRequestsSigned=true, introduce a separate SP private key and SP public certificate and sign the Redirect query correctly.", RED)

doc.add_heading("5. Application and database setup", level=1)
doc.add_heading("5.1 Fresh or cloned environment", level=2)
add_code(doc, """git clone <repository-url> cebu_rooms2
cd cebu_rooms2
composer install
npm ci
cp .env.example .env
php artisan key:generate
php artisan migrate
npm run build""")
add_p(doc, "For an existing deployment, use composer install --no-dev --optimize-autoloader and take a database backup before php artisan migrate --force.")
doc.add_heading("5.2 Database records", level=2)
add_table(doc, ["Table", "Important fields / constraints", "Operational purpose"], [
    ("saml_configurations", "unique slug/entity_id; mode; URLs; x509_cert; status; is_active; metadata; signing flags", "Federation partner configuration. Only one active record per mode is maintained by the controller."),
    ("saml_replay_records", "unique request_id, response_id, assertion_id; expires_at", "Correlates requests and blocks response/assertion replay."),
    ("saml_audit_events", "event_name, outcome, partner/user IDs, request/response IDs, IP, metadata", "Evidence and troubleshooting without storing entire assertions."),
    ("user_accounts", "email, account_status, role/permissions, last_login_at, last_login_ip", "Local authorization source; SAML authenticates but does not auto-provision."),
], [1900, 3560, 3900])
doc.add_heading("5.3 Seed strategy", level=2)
add_p(doc, "For production AMS, set the real certificate first, then run the dedicated seeder. It deactivates other IdP records and upserts the AMS record as active.")
add_code(doc, "php artisan db:seed --class=AmsSamlConfigurationSeeder --force")
add_callout(doc, "Do not use local fixtures in production", "SamlConfigurationSeeder contains a local OnePortal certificate and loopback URLs intended for development/testing. Use AmsSamlConfigurationSeeder or the admin UI with authoritative IdP metadata in production.", GOLD)

doc.add_heading("6. Environment and SAML configuration", level=1)
doc.add_heading("6.1 Environment values", level=2)
add_code(doc, """APP_URL=https://rooms.your-domain.edu.ph

SAML_SP_ENTITY_ID="${APP_URL}/saml2/metadata"
SAML_IDP_ENTITY_ID="https://ams.upcebu.edu.ph/saml2/metadata"
SAML_IDP_PUBLIC_CERT="-----BEGIN CERTIFICATE-----...-----END CERTIFICATE-----"
SAML_IDP_PRIVATE_KEY=
SAML_IDP_KEY_PASSPHRASE=

SAML_ASSERTION_TTL_SECONDS=300
SAML_CLOCK_SKEW_SECONDS=60
SAML_REQUIRE_SIGNED_REQUESTS=false
SAML_SIGN_RESPONSES=true
SAML_SIGN_ASSERTIONS=true""")
add_p(doc, "Store the certificate as a single environment value or use the platform’s supported multiline-secret mechanism. Never commit private keys or production certificates with private material. The IdP signing certificate itself is public trust material, but its lifecycle still requires controlled change management.")
doc.add_heading("6.2 Clear and rebuild Laravel caches", level=2)
add_code(doc, """php artisan config:clear
php artisan route:clear
php artisan cache:clear
php artisan view:clear
php artisan optimize""")
doc.add_heading("6.3 Configure through Rooms", level=2)
add_p(doc, "Sign in as an admin or sysadmin and open /SamlIntegration (Settings → SAML Config). Import authoritative IdP metadata XML where possible; otherwise enter the fields below manually.")
add_table(doc, ["Field", "AMS example / rule"], [
    ("Name", "UP Cebu AMS IdP"),
    ("Mode", "IdP"),
    ("Entity ID", "https://ams.upcebu.edu.ph/saml2/metadata — replace if metadata differs"),
    ("SSO URL", "https://ams.upcebu.edu.ph/saml2/sso — replace if metadata differs"),
    ("SLO URL", "Use metadata value; may remain blank until federated logout is implemented"),
    ("X.509 certificate", "Exact IdP signing certificate; no private key"),
    ("Signing algorithm", "rsa-sha256"),
    ("Default RelayState", "/MainDashboard"),
    ("Status / active", "active / true"),
    ("Attributes", "email, first_name, last_name, display_name, role, department"),
], [2400, 6960])
add_callout(doc, "Activation rule", "Creating or updating an active configuration deactivates every other provider in the same mode. Confirm the intended record before activating it.")

doc.add_heading("7. IdP registration and attribute contract", level=1)
doc.add_heading("7.1 Register the SP", level=2)
for x in [
    "Open the public metadata URL and save/import the XML into AMS/OnePortal.",
    "Verify the Entity ID and ACS URL exactly match the production APP_URL.",
    "Select HTTP-POST for SAMLResponse delivery and emailAddress for NameID.",
    "Configure signed assertions or signed responses according to the agreed profile; retain SHA-256 algorithms.",
    "Restrict the service to approved groups if the IdP supports assignment, while still enforcing Rooms local-account authorization.",
]: add_number(doc, x)
doc.add_heading("7.2 Identity mapping", level=2)
add_table(doc, ["SAML value", "Rooms behavior", "Requirement"], [
    ("NameID", "Preferred source of email.", "Email address format and value recommended."),
    ("email / mail / EmailAddress / emailAddress / email OID", "Fallback email source when NameID is empty.", "One unambiguous value matching user_accounts.email."),
    ("Other attributes", "Captured in parsed assertion data but not used to auto-provision or grant permissions.", "Treat as informational unless code is deliberately extended."),
    ("RelayState", "Accepted only when it begins with /; otherwise MainDashboard is used.", "Use relative application paths only."),
], [2500, 3550, 3310])
add_callout(doc, "Authorization boundary", "Roles and permissions continue to come from the local user_accounts record. Do not trust a SAML role attribute for authorization unless a separately reviewed mapping and entitlement lifecycle is implemented.", GOLD)
doc.add_heading("7.3 Provision test users", level=2)
add_p(doc, "Before testing, create or confirm a Rooms user account whose email equals the IdP NameID after trimming and case normalization. Set account_status to active and assign the minimum required local permissions. Also prepare an inactive user and an IdP-only user for negative tests.")

doc.add_heading("8. End-to-end test procedure", level=1)
doc.add_heading("8.1 Automated verification", level=2)
add_code(doc, """php artisan config:clear
php artisan test --filter=SamlIntegrationTest
php artisan route:list | grep -E 'saml2/(login|acs|metadata|logout)'""")
add_p(doc, "The feature suite verifies metadata secrecy, configuration import, admin authorization, single-active-provider behavior, successful ACS login, unknown-user rejection, and request correlation.")
doc.add_heading("8.2 Manual federation test", level=2)
for x in [
    "Open ${APP_URL}/saml2/metadata and confirm HTTP 200, XML content type, correct public URLs, and no private key material.",
    "Use a private browser session and open ${APP_URL}/saml2/login.",
    "Confirm redirect host, issuer, ACS, and RelayState at the IdP; authenticate with the known active test user.",
    "Confirm POST to /saml2/acs, redirect to /MainDashboard, a regenerated session, and updated last-login fields.",
    "Confirm saml.sp.request.issued and saml.sp.assertion.accepted audit events.",
    "Repeat with an unknown IdP user; expect /saml2/user-not-found and no local account creation.",
    "Repeat with an inactive Rooms user; expect the same safe failure page with inactive reason.",
    "Replay the same SAMLResponse in a controlled test environment; expect rejection.",
]: add_number(doc, x)
doc.add_heading("8.3 Acceptance matrix", level=2)
add_table(doc, ["Scenario", "Expected result"], [
    ("Valid signed response; known active email", "Login succeeds; dashboard redirect; accepted audit event."),
    ("Unknown email", "No account created; user-not-found page; rejected audit event."),
    ("Inactive local account", "No login; inactive reason recorded."),
    ("Wrong issuer/audience/destination/recipient", "ACS rejects with the corresponding validation message."),
    ("Expired/not-yet-valid assertion", "ACS rejects after applying the configured 60-second skew."),
    ("Reused response or assertion ID", "ACS rejects as already used."),
    ("Invalid or missing signature/certificate", "OneLogin validation rejects; event is audited."),
    ("External RelayState URL", "Ignored; fallback redirects to MainDashboard."),
], [3900, 5460])

doc.add_heading("9. Deployment, rollback, and operations", level=1)
doc.add_heading("9.1 Production deployment", level=2)
for x in [
    "Back up the database and current .env/secrets; record the current release identifier.",
    "Deploy code and run PHP syntax checks for the SAML and login controllers.",
    "Run composer install --no-dev --optimize-autoloader.",
    "Run php artisan migrate --force.",
    "Clear Laravel caches, seed/configure AMS, then run php artisan optimize.",
    "Verify routes and metadata, reload PHP-FPM/nginx where applicable, then perform a known-user smoke test.",
    "Monitor rejected audit events and application logs closely during the stabilization window.",
]: add_number(doc, x)
add_p(doc, "The repository script scripts/deploy-saml-fix.sh automates this sequence. Set APP_DIR, PHP_FPM_SERVICE, and WEB_SERVICE for the target host before using it; review the hard-coded final test URL.")
add_code(doc, """APP_DIR=/var/www/html \
PHP_FPM_SERVICE=php8.2-fpm \
WEB_SERVICE=nginx \
bash scripts/deploy-saml-fix.sh""")
doc.add_heading("9.2 Rollback", level=2)
for x in [
    "Deactivate the new IdP configuration or restore the previously active record through the database/admin UI.",
    "Restore the previous application release and cached configuration; retain audit records for investigation.",
    "If a migration rollback is necessary, confirm no required audit/replay/configuration data will be lost before running it.",
    "Keep local login available until federation is stable; communicate the temporary access path only to authorized administrators.",
]: add_bullet(doc, x)
doc.add_heading("9.3 Operational controls", level=2)
for x in [
    "Schedule deletion of expired saml_replay_records to control table growth.",
    "Alert on sustained saml.sp.assertion.rejected volume, certificate expiry, and metadata/issuer changes.",
    "Review SAML audit retention and personal-data handling; metadata may contain email on rejected/accepted events.",
    "Test the federation after proxy, hostname, TLS, IdP, certificate, clock, or session configuration changes.",
    "Document certificate rollover with an overlap window or dual-certificate support before the current certificate expires.",
]: add_bullet(doc, x)

doc.add_heading("10. Troubleshooting", level=1)
add_table(doc, ["Symptom / message", "Likely cause", "Action"], [
    ("SAML login is not configured", "No IdP record is both status=active and is_active=true, or SSO URL is blank.", "Activate exactly one IdP and verify its SSO URL."),
    ("Issuer does not match", "Configured Entity ID differs from Response Issuer.", "Copy the exact entityID from authoritative IdP metadata."),
    ("Destination or recipient mismatch", "APP_URL/proxy scheme/host/port or IdP ACS registration differs.", "Correct public URL generation, trusted proxy behavior, and registered ACS."),
    ("Audience mismatch", "SAML_SP_ENTITY_ID differs from Audience.", "Align the IdP audience with the exact SP Entity ID."),
    ("Response does not match active login request", "Unknown/expired InResponseTo or replay-record persistence problem.", "Start a new SP-initiated login; inspect request record and TTL."),
    ("Assertion expired / not valid yet", "Clock drift or incorrect assertion validity window.", "Synchronize clocks; keep skew small and justified."),
    ("Invalid signature or assertion", "Wrong/expired certificate, unsigned message profile, altered XML, or algorithm mismatch.", "Refresh trusted IdP certificate and confirm signing profile with IdP admin."),
    ("Email NameID missing", "IdP sends a different NameID and no recognized email attribute.", "Map NameID to emailAddress or release a supported email attribute."),
    ("User-not-found page", "Valid federation but no matching active local account.", "Provision/activate user_accounts entry; verify spelling and domain."),
    ("Redirect loop or HTTP URL in metadata", "Incorrect APP_URL or reverse proxy scheme forwarding.", "Set production HTTPS APP_URL, trusted proxies, and clear config cache."),
], [2550, 3420, 3390])
doc.add_heading("10.1 Evidence to collect", level=2)
for x in [
    "Timestamp and timezone, request ID/response ID, user email (appropriately protected), and source IP.",
    "Active saml_configurations record without exposing secrets or private keys.",
    "Relevant saml_audit_events row and Laravel log lines.",
    "Decoded SAMLResponse only in a controlled secure environment; never paste production assertions into public tools.",
    "Current metadata output, APP_URL, SAML_SP_ENTITY_ID, certificate fingerprint/expiry, and proxy headers.",
]: add_bullet(doc, x)

doc.add_heading("11. Replication checklist", level=1)
for x in [
    "[ ] Production URL, TLS, proxy, and NTP are correct.",
    "[ ] OneLogin PHP-SAML dependency and PHP XML/OpenSSL support are installed.",
    "[ ] Three SAML migrations have completed.",
    "[ ] Environment values are present and Laravel caches were rebuilt.",
    "[ ] SP metadata shows the correct Entity ID, ACS, SLO, and NameID format.",
    "[ ] IdP metadata/certificate were imported from an authoritative source.",
    "[ ] Exactly one IdP record is active.",
    "[ ] IdP has registered the exact SP Entity ID and ACS URL.",
    "[ ] Email NameID/attribute mapping is confirmed.",
    "[ ] Active, inactive, and unknown test identities are prepared.",
    "[ ] Automated SAML tests pass.",
    "[ ] Positive and negative federation tests pass with audit evidence.",
    "[ ] Rollback and non-SAML administrator access are proven.",
    "[ ] Certificate expiry, replay cleanup, logs, alerts, and ownership are documented.",
]: add_bullet(doc, x)

doc.add_heading("Appendix A · Configuration source of truth", level=1)
add_table(doc, ["Value", "Source", "Consumed by"], [
    ("SP Entity ID", "SAML_SP_ENTITY_ID, default APP_URL/saml2/metadata", "Metadata, AuthnRequest Issuer, audience validation, OneLogin settings."),
    ("ACS URL", "Laravel url('/saml2/acs')", "Metadata, AuthnRequest, recipient/destination validation."),
    ("IdP Entity ID", "Active saml_configurations.entity_id", "Issuer validation and OneLogin IdP settings."),
    ("IdP SSO/SLO", "Active saml_configurations URLs", "Redirect destination and library settings."),
    ("IdP certificate", "Active saml_configurations.x509_cert", "SAML signature validation."),
    ("TTL / skew", "SAML_ASSERTION_TTL_SECONDS / SAML_CLOCK_SKEW_SECONDS", "Request/replay expiry and assertion timing."),
], [2000, 3600, 3760])

doc.add_heading("Appendix B · Known gaps and recommended backlog", level=1)
for x in [
    "P0: Separate SP signing keys/certificates from IdP trust certificates; do not publish the IdP certificate as an SP signing key.",
    "P0: If metadata declares AuthnRequestsSigned=true, implement correct Redirect-binding signing and validate it with the IdP.",
    "P1: Align configurable signature requirements with the OneLogin security options; document whether response, assertion, or both must be signed.",
    "P1: Implement full SAML Single Logout only if required, including LogoutRequest/LogoutResponse validation and session correlation.",
    "P1: Add certificate-expiry monitoring and a rollover strategy that accepts old and new signing certificates during transition.",
    "P2: Add scheduled replay-record cleanup and operational dashboards/alerts for accepted and rejected events.",
    "P2: Add tests for wrong issuer/audience/destination, expired assertions, replay, inactive account, unsafe RelayState, and certificate rollover.",
    "P2: Consider encrypted assertions only when required by the risk assessment; key management and decryption must be designed explicitly.",
]: add_bullet(doc, x)

doc.add_heading("Appendix C · Repository reference", level=1)
add_p(doc, "This guide is based on the following implementation artifacts in the cebu_rooms2 workspace:")
for x in [
    "app/Http/Controllers/SamlSpController.php",
    "app/Http/Controllers/SamlMetadataController.php",
    "app/Http/Controllers/SamlConfigurationController.php",
    "app/Models/SamlConfiguration.php, SamlReplayRecord.php, SamlAuditEvent.php",
    "database/migrations/2026_07_06_000001 through 000003",
    "database/seeders/AmsSamlConfigurationSeeder.php and SamlConfigurationSeeder.php",
    "resources/js/Pages/SamlIntegration.vue",
    "tests/Feature/SamlIntegrationTest.php",
    "docs/saml-ams-upcebu-configuration.md",
    "scripts/deploy-saml-fix.sh",
]: add_bullet(doc, x)
add_p(doc, "End of guide.", after=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
